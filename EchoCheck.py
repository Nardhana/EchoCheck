import os
import re
import base64
import io
import json
from flask import Flask, request, jsonify
import requests
from flask_cors import CORS
from serpapi import GoogleSearch
from dotenv import load_dotenv
from google import genai
from google.genai import types

load_dotenv()

app = Flask(__name__)
CORS(app)

SERPAPI_API_KEY = os.environ.get("SERPAPI_API_KEY", "")

# Support multiple Gemini keys for rotation when one hits quota
_RAW_KEYS = [
    os.environ.get("GEMINI_API_KEY", ""),
    os.environ.get("GEMINI_API_KEY_2", ""),
]
GEMINI_KEYS = [k for k in _RAW_KEYS if k.strip()]

# Verified working models (tested 2026-04-23 via ListModels)
GEMINI_MODELS = [
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
    "gemini-2.0-flash",
    "gemini-2.0-flash-lite",
]

if not GEMINI_KEYS:
    print("[WARNING] No Gemini API keys set. Add GEMINI_API_KEY to .env")
if not SERPAPI_API_KEY:
    print("[WARNING] SERPAPI_API_KEY not set — search will be skipped, using local analysis only.")


# ─────────────────────────────────────────────
#  LOCAL FALLBACK ANALYZER
#  Runs when ALL Gemini API calls fail.
#  Uses rule-based NLP heuristics.
# ─────────────────────────────────────────────

DEBUNK_SIGNALS = [
    r"\ball\b.{0,30}\b(politicians?|governments?|banks?|scientists?)\b.{0,30}\b(corrupt|lie|lying|fake|fraud)",
    r"\b(proven|confirmed|exposed)\b.{0,40}\b(hoax|conspiracy|scam|cover.?up)",
    r"\b(they|the government|elites?|illuminati|deep state)\b.{0,40}\b(want|hiding|conceal|suppress)",
    r"\bwake up\b",
    r"\b(miracle|instant|guaranteed)\b.{0,20}\bcure\b",
    r"\b100%\b.{0,20}\b(safe|effective|proven|guaranteed)\b",
    r"\bcauses?\b.{0,20}\b(autism|cancer|death)\b.{0,20}\b(vaccine|5g|wifi|fluoride)",
    r"\b(never happened|didn.?t happen|staged|crisis actor)",
    r"\bflat earth\b",
    r"\bearth is flat\b",
    r"\bsun rises.{0,10}west\b",
]

CONFIRM_SIGNALS = [
    r"\b(according to|reported by|study shows|research shows|published in)\b",
    r"\b(official|government|who|cdc|nasa|fbi|cia|un|eu)\b.{0,30}\b(report|statement|confirms?|announced?)\b",
    r"\b(percent|%|million|billion|trillion)\b",
    r"\b\d{4}\b",
    r"\b(passed|signed|enacted|introduced)\b.{0,30}\b(law|bill|act|legislation)\b",
]

COMPLEX_SIGNALS = [
    r"\b(some|many|often|sometimes|can|may|might|debate|controversial|disputed)\b",
    r"\b(both|however|although|while|despite|but|yet)\b",
    r"\b(depends|context|situation|case by case)\b",
]

MISINFORMATION_KEYWORDS = {
    "plandemic", "scamdemic", "bioweapon", "chemtrails", "microchip",
    "nwo", "new world order", "false flag", "crisis actor",
    "reptilian", "qanon", "moon landing fake",
}

# ─────────────────────────────────────────────
#  DOMAIN KNOWLEDGE BASE
#  Based on AllSides Media Bias Ratings +
#  MediaBiasFactCheck credibility ratings.
#  Domain-map values override Gemini estimates
#  (more consistent and defensible).
# ─────────────────────────────────────────────

DOMAIN_BIAS_MAP = {
    # Left-leaning
    "nytimes.com": "Left-leaning",
    "washingtonpost.com": "Left-leaning",
    "theguardian.com": "Left-leaning",
    "huffpost.com": "Left-leaning",
    "vox.com": "Left-leaning",
    "msnbc.com": "Left-leaning",
    "slate.com": "Left-leaning",
    "motherjones.com": "Left-leaning",
    "thedailybeast.com": "Left-leaning",
    "theatlantic.com": "Left-leaning",
    "salon.com": "Left-leaning",
    # Left-Center
    "cnn.com": "Left-Center",
    "npr.org": "Left-Center",
    "abcnews.go.com": "Left-Center",
    "nbcnews.com": "Left-Center",
    "cbsnews.com": "Left-Center",
    "time.com": "Left-Center",
    "politico.com": "Left-Center",
    "buzzfeednews.com": "Left-Center",
    # Center
    "reuters.com": "Center",
    "apnews.com": "Center",
    "bbc.com": "Center",
    "bbc.co.uk": "Center",
    "usatoday.com": "Center",
    "thehill.com": "Center",
    "pbs.org": "Center",
    "axios.com": "Center",
    "politifact.com": "Center",
    "snopes.com": "Center",
    "factcheck.org": "Center",
    "wikipedia.org": "Center",
    "nature.com": "Center",
    "science.org": "Center",
    "nih.gov": "Center",
    "cdc.gov": "Center",
    "who.int": "Center",
    "nasa.gov": "Center",
    "economist.com": "Center",
    "csmonitor.com": "Center",
    # Right-Center
    "wsj.com": "Right-Center",
    "forbes.com": "Right-Center",
    "reason.com": "Right-Center",
    # Right-leaning
    "foxnews.com": "Right-leaning",
    "nypost.com": "Right-leaning",
    "nationalreview.com": "Right-leaning",
    "washingtonexaminer.com": "Right-leaning",
    "newsmax.com": "Right-leaning",
    "theblaze.com": "Right-leaning",
    "westernjournal.com": "Right-leaning",
    "epochtimes.com": "Right-leaning",
    "breitbart.com": "Right-leaning",
    "dailywire.com": "Right-leaning",
    "oann.com": "Right-leaning",
    "infowars.com": "Right-leaning",
}

DOMAIN_CREDIBILITY_MAP = {
    # High — peer-reviewed, government agencies, established wire services
    "reuters.com": "High", "apnews.com": "High",
    "bbc.com": "High", "bbc.co.uk": "High",
    "pbs.org": "High", "nature.com": "High",
    "science.org": "High", "nih.gov": "High",
    "cdc.gov": "High", "who.int": "High",
    "nasa.gov": "High", "factcheck.org": "High",
    "snopes.com": "High", "politifact.com": "High",
    "scholar.google.com": "High", "pubmed.ncbi.nlm.nih.gov": "High",
    # Medium — established mainstream outlets, generally reliable
    "nytimes.com": "Medium", "washingtonpost.com": "Medium",
    "theguardian.com": "Medium", "cnn.com": "Medium",
    "foxnews.com": "Medium", "npr.org": "Medium",
    "abcnews.go.com": "Medium", "nbcnews.com": "Medium",
    "cbsnews.com": "Medium", "usatoday.com": "Medium",
    "thehill.com": "Medium", "time.com": "Medium",
    "politico.com": "Medium", "wsj.com": "Medium",
    "axios.com": "Medium", "theatlantic.com": "Medium",
    "vox.com": "Medium", "nationalreview.com": "Medium",
    "washingtonexaminer.com": "Medium", "forbes.com": "Medium",
    "wikipedia.org": "Medium", "nypost.com": "Medium",
    "newsmax.com": "Medium", "theblaze.com": "Medium",
    "dailywire.com": "Medium", "epochtimes.com": "Medium",
    "westernjournal.com": "Medium",
    # Low — known for poor sourcing or partisan fabrication
    "breitbart.com": "Low", "infowars.com": "Low",
    "thegatewaypundit.com": "Low", "naturalnews.com": "Low",
    "beforeitsnews.com": "Low", "oann.com": "Low",
}


def get_domain_info(url: str) -> dict:
    """Resolve domain from URL and return bias + credibility from knowledge base."""
    from urllib.parse import urlparse
    try:
        domain = urlparse(url).netloc.lower().lstrip("www.")
        bias = DOMAIN_BIAS_MAP.get(domain)
        credibility = DOMAIN_CREDIBILITY_MAP.get(domain)
        if not bias:
            # Try base domain for subdomains (e.g. news.bbc.com → bbc.com)
            parts = domain.split(".")
            if len(parts) >= 2:
                base = ".".join(parts[-2:])
                bias = DOMAIN_BIAS_MAP.get(base)
                credibility = DOMAIN_CREDIBILITY_MAP.get(base)
        return {"bias": bias, "credibility": credibility or "Unknown"}
    except Exception:
        return {"bias": None, "credibility": "Unknown"}


def enrich_evidence_items(evidence: list, search_results: list) -> list:
    """
    Attach URLs from SerpAPI results and override bias/credibility
    using the domain knowledge base (more reliable than free-form LLM estimates).
    """
    url_map = {r.get("title", ""): r.get("link", "#") for r in search_results}
    for item in evidence:
        if not item.get("url") or item.get("url") == "#":
            item["url"] = url_map.get(item.get("title", ""), "#")
        domain_info = get_domain_info(item.get("url", "#"))
        if domain_info["bias"]:
            item["bias"] = domain_info["bias"]
        elif not item.get("bias"):
            item["bias"] = "Center"
        item["credibility"] = domain_info["credibility"]
    return evidence


def local_fallback_analysis(statement: str, search_results: list) -> dict:
    """
    Rule-based fallback analyzer. Runs when Gemini is unavailable.
    Uses regex heuristics + keyword matching + search snippet analysis.
    """
    print("\n[FALLBACK] Running local rule-based analysis...")
    text = statement.lower()
    words = set(re.findall(r'\b\w+\b', text))

    debunk_hits = sum(1 for p in DEBUNK_SIGNALS if re.search(p, text, re.IGNORECASE))
    confirm_hits = sum(1 for p in CONFIRM_SIGNALS if re.search(p, text, re.IGNORECASE))
    complex_hits = sum(1 for p in COMPLEX_SIGNALS if re.search(p, text, re.IGNORECASE))
    misinfo_hits = len(words & MISINFORMATION_KEYWORDS)

    snippet_text = " ".join(
        (r.get("snippet", "") + " " + r.get("title", "")).lower()
        for r in search_results[:5]
    )
    snippet_debunk = sum(1 for p in DEBUNK_SIGNALS if re.search(p, snippet_text, re.IGNORECASE))
    snippet_confirm = sum(1 for p in CONFIRM_SIGNALS if re.search(p, snippet_text, re.IGNORECASE))

    total_debunk = debunk_hits + misinfo_hits * 2 + snippet_debunk
    total_confirm = confirm_hits + snippet_confirm

    if total_debunk >= 2 or misinfo_hits >= 1:
        verdict = "Debunked"
        confidence = min(40 + total_debunk * 8 + misinfo_hits * 15, 72)
        reasoning = (
            "Local analysis detected language patterns and keywords commonly associated with "
            "misinformation or debunked claims. This is an automated heuristic result — "
            "AI verification was unavailable. Treat with caution and verify independently."
        )
    elif complex_hits >= 2 and total_confirm == 0:
        verdict = "Complex/Mixed"
        confidence = 30
        reasoning = (
            "The statement contains qualifying language suggesting it may be context-dependent "
            "or partially true. AI verification was unavailable; this is a heuristic estimate. "
            "Manual verification is strongly recommended."
        )
    elif total_confirm >= 2:
        verdict = "Inconclusive"
        confidence = 35
        reasoning = (
            "The statement contains factual language markers, but AI verification was unavailable "
            "to confirm against live sources. Please retry when the AI service is available."
        )
    else:
        verdict = "Inconclusive"
        confidence = 20
        reasoning = (
            "AI analysis service is currently unavailable (API quota/connectivity issue). "
            "A local heuristic analysis could not determine the validity of this claim. "
            "Please retry in a few minutes or check the server logs."
        )

    evidence = []
    for r in search_results[:3]:
        evidence.append({
            "title": r.get("title", "Unknown"),
            "source": r.get("source", r.get("displayed_link", "Unknown")),
            "snippet": r.get("snippet", "No snippet available."),
            "bias": "Center",
            "url": r.get("link", "#"),
        })
    evidence = enrich_evidence_items(evidence, search_results)

    return {
        "verdict": verdict,
        "reasoning": reasoning,
        "confidence": confidence,
        "evidence": evidence,
        "fallback_used": True,
    }


# ─────────────────────────────────────────────
#  GEMINI CALL  (multi-key rotation via SDK)
# ─────────────────────────────────────────────

def _call_gemini_sdk(prompt_text: str, image_data: dict = None):
    """
    Tries every key × every model. Returns raw response text or None.
    image_data: {'bytes': bytes, 'mime_type': str} for vision requests.
    """
    last_error = ""
    for key in GEMINI_KEYS:
        client = genai.Client(api_key=key)
        for model in GEMINI_MODELS:
            try:
                if image_data:
                    if "2.5" not in model:
                        continue
                    part_image = types.Part.from_bytes(
                        data=image_data["bytes"],
                        mime_type=image_data["mime_type"],
                    )
                    contents = [part_image, prompt_text]
                else:
                    contents = prompt_text

                response = client.models.generate_content(
                    model=model,
                    contents=contents,
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                    ),
                )
                print(f"--> Success: key=...{key[-6:]} model={model}")
                return response.text
            except Exception as e:
                msg = str(e)
                code = "429" if "429" in msg else "403" if "403" in msg else "ERR"
                print(f"--> FAIL key=...{key[-6:]} model={model}: [{code}] {msg[:100]}")
                last_error = msg
    print(f"[ERROR] All Gemini keys/models exhausted. Last: {last_error[:200]}")
    return None


# ─────────────────────────────────────────────
#  CORE ANALYSIS
# ─────────────────────────────────────────────

def fetch_google_search_results(query: str) -> list:
    print("\nPerforming real-time Google Search...")
    if not SERPAPI_API_KEY:
        print("--> SerpApi key not set. Skipping search.")
        return []
    try:
        search = GoogleSearch({"q": query, "api_key": SERPAPI_API_KEY})
        results = search.get_dict().get("organic_results", [])
        print(f"--> Found {len(results)} results.")
        return results
    except Exception as e:
        print(f"--> Search failed: {e}")
        return []


def get_ai_analysis(statement: str, search_results: list) -> dict:
    print("\nSending to Gemini AI...")

    evidence_text = "\n".join(
        f"Source: {r.get('source','N/A')}\nTitle: {r.get('title','N/A')}\nSnippet: {r.get('snippet','N/A')}\n"
        for r in search_results[:5]
    ) or "No search results available."

    prompt = f"""You are EchoCheck, an expert AI fact-checker. Analyze the statement using chain-of-thought reasoning.

Statement: "{statement}"

Real-time Search Evidence:
---
{evidence_text}
---

Steps:
1. Identify core factual claims.
2. Match each claim against the evidence ONLY — do NOT use your own training knowledge or assume facts not present in the evidence.
3. Verdict: "Confirmed", "Debunked", "Complex/Mixed", or "Inconclusive".
4. Confidence: integer 0-100, reflecting how strongly the evidence supports the verdict.
5. Select the 3 most relevant evidence items. For each, estimate political bias as exactly one of: "Left-leaning", "Left-Center", "Center", "Right-Center", "Right-leaning".

Respond ONLY as valid JSON with no markdown fencing:
{{
  "verdict": string,
  "reasoning": string (2-4 sentences, cite specific evidence titles),
  "confidence": integer,
  "evidence": [
    {{"title": string, "source": string, "snippet": string, "bias": string}},
    ...
  ]
}}"""

    raw = _call_gemini_sdk(prompt)

    if raw is None:
        print("[FALLBACK] Gemini unavailable — using local analyzer.")
        return local_fallback_analysis(statement, search_results)

    try:
        clean = raw.strip()
        if clean.startswith("```"):
            clean = re.sub(r"^```[a-z]*\n?", "", clean).rstrip("```").strip()
        data = json.loads(clean)
        print(f"--> Verdict: {data.get('verdict')} ({data.get('confidence','?')}%)")
        data["evidence"] = enrich_evidence_items(data.get("evidence", []), search_results)
        data["fallback_used"] = False
        return data
    except Exception as e:
        print(f"--> Failed to parse Gemini response: {e}\nRaw: {raw[:300]}")
        return local_fallback_analysis(statement, search_results)


def extract_text_from_document(file_bytes: bytes, filename: str):
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "pdf":
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            return "\n".join(p.extract_text() or "" for p in reader.pages).strip()
        elif ext == "docx":
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs).strip()
        elif ext == "txt":
            return file_bytes.decode("utf-8", errors="ignore").strip()
    except Exception as e:
        print(f"--> Document extraction failed: {e}")
    return None


def analyze_image_with_gemini(image_bytes: bytes, mime_type: str, statement: str = "") -> dict:
    print("\nAnalyzing image with Gemini Vision...")
    claim_line = f'\nClaim about this image: "{statement}"' if statement else ""
    prompt = f"""You are an AI media forensics analyst. Scrutinize this image for:
1. Manipulation Indicators: editing artifacts, cloning, splicing, inconsistent lighting/shadows.
2. AI/Deepfake Signs: unnatural facial features, texture anomalies, over-smooth skin.
3. Contextual Accuracy: does the visual content match what is claimed?{claim_line}

Respond ONLY as valid JSON:
{{
  "verdict": "Authentic"|"Likely Manipulated"|"AI Generated"|"Deepfake Suspected"|"Inconclusive",
  "manipulation_score": integer 0-100,
  "confidence": integer 0-100,
  "reasoning": string,
  "indicators": [string, ...],
  "claim_match": "Supports"|"Contradicts"|"Unrelated"|"Cannot Determine"|null
}}"""

    raw = _call_gemini_sdk(prompt, image_data={"bytes": image_bytes, "mime_type": mime_type})

    fallback_img = {
        "verdict": "Inconclusive",
        "manipulation_score": -1,
        "confidence": 0,
        "reasoning": "AI analysis unavailable. Image could not be assessed automatically.",
        "indicators": [],
        "claim_match": None,
        "fallback_used": True,
    }

    if raw is None:
        return fallback_img
    try:
        clean = raw.strip()
        if clean.startswith("```"):
            clean = re.sub(r"^```[a-z]*\n?", "", clean).rstrip("```").strip()
        result = json.loads(clean)
        result["fallback_used"] = False
        return result
    except Exception as e:
        print(f"--> Failed to parse image response: {e}")
        return fallback_img


# ─────────────────────────────────────────────
#  ROUTES
# ─────────────────────────────────────────────

@app.route("/", methods=["GET"])
def index():
    return jsonify({"status": "ok", "message": "EchoCheck Server is running."})


@app.route("/health", methods=["GET"])
def health():
    """System health check — useful for demo reliability verification."""
    return jsonify({
        "status": "ok",
        "version": "2.0.0",
        "gemini_keys_loaded": len(GEMINI_KEYS),
        "search_enabled": bool(SERPAPI_API_KEY),
        "active_models": GEMINI_MODELS,
        "domain_knowledge_base": {
            "bias_entries": len(DOMAIN_BIAS_MAP),
            "credibility_entries": len(DOMAIN_CREDIBILITY_MAP),
        },
        "fallback_available": True,
    })


@app.route("/analyze", methods=["POST"])
def analyze_claim():
    data = request.get_json()
    if not data or "statement" not in data:
        return jsonify({"error": 'Invalid request. "statement" key is required.'}), 400
    statement = data["statement"].strip()
    print(f"\n\n{'='*50}\nQuery: {statement}")
    search_results = fetch_google_search_results(statement)
    return jsonify(get_ai_analysis(statement, search_results))


@app.route("/analyze-image", methods=["POST"])
def analyze_image():
    if "image" not in request.files:
        return jsonify({"error": "No image provided."}), 400
    file = request.files["image"]
    statement = request.form.get("statement", "").strip()
    if file.content_type not in {"image/jpeg", "image/png", "image/webp", "image/gif"}:
        return jsonify({"error": "Unsupported type. Use JPEG, PNG, WebP, or GIF."}), 400
    image_bytes = file.read()
    if len(image_bytes) > 10 * 1024 * 1024:
        return jsonify({"error": "Image exceeds 10 MB limit."}), 400
    return jsonify(analyze_image_with_gemini(image_bytes, file.content_type, statement))


@app.route("/analyze-document", methods=["POST"])
def analyze_document():
    if "document" not in request.files:
        return jsonify({"error": "No document provided."}), 400
    file = request.files["document"]
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in {"pdf", "docx", "txt"}:
        return jsonify({"error": "Unsupported type. Use PDF, DOCX, or TXT."}), 400
    file_bytes = file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        return jsonify({"error": "File exceeds 20 MB limit."}), 400
    extracted = extract_text_from_document(file_bytes, file.filename)
    if not extracted:
        return jsonify({"error": "Could not extract text from document."}), 422
    analysis_text = extracted[:3000] + ("...[truncated]" if len(extracted) > 3000 else "")
    search_results = fetch_google_search_results(" ".join(extracted.split()[:20]))
    result = get_ai_analysis(analysis_text, search_results)
    result["extracted_preview"] = extracted[:400] + ("..." if len(extracted) > 400 else "")
    result["mode"] = "document"
    return jsonify(result)


if __name__ == "__main__":
    app.run(debug=True, port=5000)
import os
import base64
import io
from flask import Flask, request, jsonify
import requests
from flask_cors import CORS
import json
from serpapi import GoogleSearch
from dotenv import load_dotenv

load_dotenv()  # loads from .env file in project root

app = Flask(__name__)
CORS(app)

SERPAPI_API_KEY = os.environ.get("SERPAPI_API_KEY", "")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")

if not GEMINI_API_KEY:
    print("[WARNING] GEMINI_API_KEY is not set. Add it to your .env file.")
if not SERPAPI_API_KEY:
    print("[WARNING] SERPAPI_API_KEY is not set. Add it to your .env file.")


@app.route('/', methods=['GET'])
def index():
    return jsonify({'status': 'ok', 'message': 'EchoCheck RAG Server is running.'})

def perform_sanity_check(statement):
    lower_case_statement = statement.lower()
    impossible_claims = [
        {'keywords': ['sun', 'rises', 'west'], 'reason': 'This claim contradicts fundamental laws of astronomy.'},
        {'keywords': ['earth', 'flat'], 'reason': 'This claim contradicts centuries of established scientific evidence.'}
    ]
    for claim in impossible_claims:
        if all(kw in lower_case_statement for kw in claim['keywords']):
            return {'passed': False, 'reason': claim['reason']}
    return {'passed': True, 'reason': None}

def fetch_google_search_results(query):
    """
    Performs a real-time Google search to get the latest information.
    """
    print("\nPerforming real-time Google Search...")
    if not SERPAPI_API_KEY:
        print("--> SerpApi key not set.")
        return []
    try:
        params = {
            "q": query,
            "api_key": SERPAPI_API_KEY
        }
        search = GoogleSearch(params)
        results = search.get_dict()
        organic_results = results.get("organic_results", [])
        print(f"--> Found {len(organic_results)} search results.")
        return organic_results
    except Exception as e:
        print(f"--> FAILED to fetch Google Search results: {e}")
        return []

def get_ai_analysis(statement, search_results):
    print("\nSending statement and search results to Gemini AI...")
    if not GEMINI_API_KEY:
        return {'verdict': 'API Error', 'reasoning': 'Gemini API key is not set.', 'evidence': [], 'confidence': 0}

    evidence_snippets = []
    for result in search_results[:5]:
        snippet = f"Source: {result.get('source', 'N/A')}\nTitle: {result.get('title')}\nSnippet: {result.get('snippet', 'N/A')}\n"
        evidence_snippets.append(snippet)
    
    evidence_text = "\n".join(evidence_snippets)

    prompt = f"""
    You are EchoCheck, an expert AI fact-checker. Analyze the following statement using structured chain-of-thought reasoning.

    Statement to analyze: "{statement}"

    Real-time Search Evidence:
    ---
    {evidence_text}
    ---

    Think step by step:
    1. Identify the core factual claims in the statement.
    2. Match each claim against the provided evidence — do NOT rely on your own training knowledge.
    3. Weigh the overall evidence and determine a verdict: "Confirmed", "Debunked", "Complex/Mixed", or "Inconclusive" (if evidence is insufficient).
    4. Rate your confidence in the verdict as an integer from 0 to 100.
    5. Select the 3 most relevant evidence items. For each, estimate its political bias as "Left-leaning", "Center", or "Right-leaning".

    Respond in a single, strict JSON format with four keys:
    - "verdict": string
    - "reasoning": string (2-3 sentences referencing the evidence)
    - "confidence": integer 0-100
    - "evidence": JSON array of exactly 3 objects, each with keys: "title", "source", "snippet", "bias"
    """

    data = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"responseMimeType": "application/json"}
    }

    result = _call_gemini(data)
    if result is None:
        return {
            'verdict': 'API Error',
            'reasoning': 'Gemini API call failed. Check the terminal for the exact error (quota, leaked key, or missing key). Get a new key at https://aistudio.google.com/',
            'evidence': [],
            'confidence': 0
        }

    try:
        content = result['candidates'][0]['content']['parts'][0]['text']
        clean_json_string = content.strip().replace('```json', '').replace('```', '')
        
        verdict_data = json.loads(clean_json_string)
        print(f"--> Gemini Verdict: {verdict_data.get('verdict')} (confidence: {verdict_data.get('confidence', 'N/A')}%)")
        
        # Match evidence items back to search result URLs by title
        search_url_map = {r.get('title', ''): r.get('link', '#') for r in search_results}
        for item in verdict_data.get('evidence', []):
            item['url'] = search_url_map.get(item.get('title', ''), '#')

        return verdict_data

    except Exception as e:
        print(f"--> FAILED to parse Gemini response: {e}")
        return {'verdict': 'API Error', 'reasoning': 'Could not process the response from the AI model.', 'evidence': [], 'confidence': 0}


# Ordered list: (api_version, model_name). Verified via ListModels on 2026-04-23.
GEMINI_MODELS = [
    ("v1beta", "gemini-2.5-flash"),
    ("v1beta", "gemini-2.5-flash-lite"),
    ("v1beta", "gemini-2.0-flash"),
    ("v1beta", "gemini-2.0-flash-lite"),
    ("v1beta", "gemini-2.0-flash-001"),
    ("v1beta", "gemini-2.0-flash-lite-001"),
]


def _call_gemini(data, models=None):
    """Shared helper to call Gemini models with fallback. Logs real error bodies."""
    if not GEMINI_API_KEY:
        print("[ERROR] GEMINI_API_KEY is not set.")
        return None
    model_list = models if models is not None else GEMINI_MODELS
    headers = {'Content-Type': 'application/json'}
    last_error = ""
    for api_version, model in model_list:
        url = f"https://generativelanguage.googleapis.com/{api_version}/models/{model}:generateContent?key={GEMINI_API_KEY}"
        try:
            response = requests.post(url, headers=headers, json=data, timeout=45)
            if response.status_code == 429:
                err_msg = response.json().get('error', {}).get('message', response.text[:200])
                print(f"--> [{model}] 429 Quota exceeded: {err_msg}")
                last_error = f"Quota exceeded on {model}"
                continue
            if response.status_code == 403:
                err_msg = response.json().get('error', {}).get('message', response.text[:200])
                print(f"--> [{model}] 403 Forbidden: {err_msg}")
                last_error = f"403 on {model}: {err_msg}"
                continue
            if not response.ok:
                err_msg = response.text[:300]
                print(f"--> [{model}] HTTP {response.status_code}: {err_msg}")
                last_error = f"HTTP {response.status_code} on {model}"
                continue
            result = response.json()
            print(f"--> Success with model: {model} (API {api_version})")
            return result
        except requests.exceptions.Timeout:
            print(f"--> [{model}] Request timed out.")
            last_error = f"Timeout on {model}"
        except Exception as e:
            print(f"--> [{model}] Unexpected error: {e}")
            last_error = str(e)
    print(f"[ERROR] All Gemini models failed. Last error: {last_error}")
    return None


def extract_text_from_document(file_bytes, filename):
    """Extract plain text from PDF, DOCX, or TXT files."""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    try:
        if ext == 'pdf':
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            text = "\n".join(page.extract_text() or '' for page in reader.pages)
            return text.strip()
        elif ext in ('docx',):
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join(para.text for para in doc.paragraphs)
            return text.strip()
        elif ext == 'txt':
            return file_bytes.decode('utf-8', errors='ignore').strip()
    except Exception as e:
        print(f"--> Failed to extract text from document: {e}")
    return None


def analyze_image_with_gemini(image_bytes, mime_type, statement=''):
    """Use Gemini Vision to detect manipulation, deepfakes, and AI generation."""
    print("\nAnalyzing image with Gemini Vision...")
    image_b64 = base64.b64encode(image_bytes).decode('utf-8')
    claim_line = f'\n    Claim about this image: "{statement}"' if statement else ''
    prompt = f"""You are an AI media forensics analyst. Scrutinize this image for:
    1. Manipulation Indicators: editing artifacts, cloning, splicing, inconsistent lighting/shadows, unnatural edges.
    2. AI/Deepfake Signs: unnatural facial features, texture anomalies, background inconsistencies, over-smooth skin.
    3. Contextual Accuracy: does the visual content match what is claimed?{claim_line}

    Respond in strict JSON with these keys:
    - "verdict": one of "Authentic", "Likely Manipulated", "AI Generated", "Deepfake Suspected", "Inconclusive"
    - "manipulation_score": integer 0-100 (0=clearly authentic, 100=clearly fake)
    - "confidence": integer 0-100 (your confidence in this assessment)
    - "reasoning": string (2-3 sentences explaining your findings)
    - "indicators": array of strings (specific signs detected, empty array if none)
    - "claim_match": string — only if a claim was provided, else null. Values: "Supports", "Contradicts", "Unrelated", "Cannot Determine"
    """
    data = {
        "contents": [{"parts": [
            {"inline_data": {"mime_type": mime_type, "data": image_b64}},
            {"text": prompt}
        ]}],
        "generationConfig": {"responseMimeType": "application/json"}
    }
    raw = _call_gemini(data)
    if not raw:
        return {"verdict": "Error", "manipulation_score": -1, "confidence": 0,
                "reasoning": "All Gemini models failed.", "indicators": [], "claim_match": None}
    try:
        text = raw['candidates'][0]['content']['parts'][0]['text']
        return json.loads(text.strip().replace('```json', '').replace('```', ''))
    except Exception as e:
        print(f"--> Failed to parse image analysis response: {e}")
        return {"verdict": "Error", "manipulation_score": -1, "confidence": 0,
                "reasoning": "Could not parse AI response.", "indicators": [], "claim_match": None}


@app.route('/analyze-image', methods=['POST'])
def analyze_image():
    if 'image' not in request.files:
        return jsonify({'error': 'No image provided.'}), 400
    file = request.files['image']
    statement = request.form.get('statement', '').strip()
    allowed_types = {'image/jpeg', 'image/png', 'image/webp', 'image/gif'}
    if file.content_type not in allowed_types:
        return jsonify({'error': 'Unsupported type. Use JPEG, PNG, WebP, or GIF.'}), 400
    image_bytes = file.read()
    if len(image_bytes) > 10 * 1024 * 1024:
        return jsonify({'error': 'Image exceeds 10 MB limit.'}), 400
    result = analyze_image_with_gemini(image_bytes, file.content_type, statement)
    return jsonify(result)


@app.route('/analyze-document', methods=['POST'])
def analyze_document():
    if 'document' not in request.files:
        return jsonify({'error': 'No document provided.'}), 400
    file = request.files['document']
    allowed_ext = {'pdf', 'docx', 'txt'}
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in allowed_ext:
        return jsonify({'error': 'Unsupported type. Use PDF, DOCX, or TXT.'}), 400
    file_bytes = file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        return jsonify({'error': 'File exceeds 20 MB limit.'}), 400
    extracted_text = extract_text_from_document(file_bytes, file.filename)
    if not extracted_text:
        return jsonify({'error': 'Could not extract text from document.'}), 422
    # Use first 3000 chars for analysis to stay within token limits
    analysis_text = extracted_text[:3000] + ('...[truncated]' if len(extracted_text) > 3000 else '')
    print(f"\n\n--- Document Analysis Request ---\nFile: {file.filename} | Chars: {len(extracted_text)}")
    search_query = ' '.join(extracted_text.split()[:20])  # first 20 words as search query
    search_results = fetch_google_search_results(search_query)
    result = get_ai_analysis(analysis_text, search_results)
    result['extracted_preview'] = extracted_text[:400] + ('...' if len(extracted_text) > 400 else '')
    result['mode'] = 'document'
    return jsonify(result)


@app.route('/analyze', methods=['POST'])
def analyze_claim():
    data = request.get_json()
    if not data or 'statement' not in data:
        return jsonify({'error': 'Invalid request. "statement" key is required.'}), 400

    statement = data['statement'].strip()
    print(f"\n\n--- New Request Received ---\nQuery: {statement}")

    sanity_check = perform_sanity_check(statement)
    if not sanity_check['passed']:
        print("--> Sanity check failed.")
        return jsonify({'verdict': 'Fundamentally False', 'reasoning': sanity_check['reason'], 'evidence': []})
    
    print("--> Sanity check passed. Fetching real-time search results...")
    
    search_results = fetch_google_search_results(statement)
    
    if not search_results:
        return jsonify({'verdict': 'Inconclusive', 'reasoning': 'Could not find any relevant information in a real-time search.', 'evidence': []}), 200

    result = get_ai_analysis(statement, search_results)
    
    return jsonify(result)

if __name__ == '__main__':
    app.run(debug=True, port=5000)
